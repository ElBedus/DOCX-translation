import os
import json
import ctypes

import docx
from googletrans import Translator


MAIN_FOLDER = os.path.dirname(__file__)
file = os.path.join(MAIN_FOLDER, "language_codes.json")
with open(file, "r") as read_file:
    LANGCODES = json.load(read_file)

kernel32 = ctypes.WinDLL('kernel32')
hStdOut = kernel32.GetStdHandle(-11)
mode = ctypes.c_ulong()
kernel32.GetConsoleMode(hStdOut, ctypes.byref(mode))
mode.value |= 4
kernel32.SetConsoleMode(hStdOut, mode)


class bcolors:
    GREEN = '\033[92m'
    CYAN = '\033[96m'
    WARNING = '\033[93m'
    FAIL = '\033[91m'
    ENDC = '\033[0m'


def input_values() -> list:
    print("\n"*2)
    source = input(f"{bcolors.CYAN}Source file (path or name):{bcolors.ENDC} ").strip()
    source = source.replace("\\", "/")
    sourceLanguageCode = input(f"{bcolors.CYAN}Source file language:{bcolors.ENDC} ").strip()
    target = input(f"{bcolors.CYAN}Target file (path or name):{bcolors.ENDC} ").strip()
    targetLanguageCode = input(f"{bcolors.CYAN}Target file language:{bcolors.ENDC} ").strip()
    print()

    return [source, sourceLanguageCode, target, targetLanguageCode]

def input_check(inputValues) -> bool:
    source, sourceLanguageCode, target, targetLanguageCode = inputValues

    if not os.path.exists(source):
        source = os.path.join(MAIN_FOLDER, source)
        if not os.path.exists(source):
            print(f"{bcolors.FAIL}Source file/path {source} does not exist.{bcolors.ENDC}")
            return False

    if sourceLanguageCode not in LANGCODES:
        print(f"{bcolors.FAIL}Invalid source file language.{bcolors.ENDC}")
        return False
    if targetLanguageCode not in LANGCODES:
        print(f"{bcolors.FAIL}Invalid target file language.{bcolors.ENDC}")
        return False
    
    return True

def docx_replace():
    k=-1
    for p in paragraphs:
        k+=1
        key_name, val = paragraphs[k].text, translatedText[k]

        if key_name in str(p.text):
            inline = p.runs
            started = False
            key_index = 0

            # found_runs is a list of (inline index, index of match, length of match)
            found_runs = list()
            found_all = False
            replace_done = False

            for i in range(len(inline)):

                # case 1: found in single run so short circuit the replace
                if key_name in inline[i].text and not started:
                    found_runs.append((i, inline[i].text.find(key_name), len(key_name)))
                    text = inline[i].text.replace(key_name, str(val))
                    inline[i].text = text
                    replace_done = True
                    found_all = True
                    break

                if key_name[key_index] not in inline[i].text and not started:
                    # keep looking ...
                    continue

                # case 2: search for partial text, find first run
                if key_name[key_index] in inline[i].text and inline[i].text[-1] in key_name and not started:
                    # check sequence
                    start_index = inline[i].text.find(key_name[key_index])
                    check_length = len(inline[i].text)
                    for text_index in range(start_index, check_length):
                        if inline[i].text[text_index] != key_name[key_index]:
                            # no match so must be false positive
                            break
                    if key_index == 0:
                        started = True
                    chars_found = check_length - start_index
                    key_index += chars_found
                    found_runs.append((i, start_index, chars_found))
                    if key_index != len(key_name):
                        continue
                    else:
                        # found all chars in key_name
                        found_all = True
                        break

                # case 3: search for partial text, find subsequent run
                if key_name[key_index] in inline[i].text and started and not found_all:
                    # check sequence
                    chars_found = 0
                    check_length = len(inline[i].text)
                    for text_index in range(0, check_length):
                        if inline[i].text[text_index] == key_name[key_index]:
                            key_index += 1
                            chars_found += 1
                        else: break
                    # no match so must be end
                    found_runs.append((i, 0, chars_found))
                    if key_index == len(key_name):
                        found_all = True
                        break

            if found_all and not replace_done:
                for i, item in enumerate(found_runs):
                    index, start, length = [t for t in item]
                    if i == 0:
                        text = inline[index].text.replace(inline[index].text[start:start + length], str(val))
                        inline[index].text = text
                    else:
                        text = inline[index].text.replace(inline[index].text[start:start + length], '')
                        inline[index].text = text

def text_translation() -> list:
    translatedText=[]
    for para in paragraphs:
        try:
            if para.text == "":
                translatedText.append("")
            else:
                translation = translator.translate(para.text,src=sourceLanguageCode,dest=targetLanguageCode)
                translatedText.append(translation.text)
        except:
            translatedText.append("")
    
    return translatedText


for key in LANGCODES:
    print(f"{LANGCODES[key]}: {key}")

while True:
    inputValues = input_values()
    success = input_check(inputValues)
    if not success: continue

    source, sourceLanguageCode, target, targetLanguageCode = inputValues

    doc = docx.Document(source)

    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraphs.append(paragraph) #C:\Users\elbed\Desktop

    translator = Translator()

    translatedText = text_translation()

    docx_replace()
    
    try: doc.save(target)
    except Exception as e:
        print(f"{bcolors.FAIL}Invalid target file/path: {source}{bcolors.ENDC}")
        print(f"{bcolors.FAIL}Error: {e}{bcolors.ENDC}")
        continue

    print(f"{bcolors.GREEN}Document translation is completed.{bcolors.ENDC}")